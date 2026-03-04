import os
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import docx

def process_document(file_path):
    # 1. Identify file type and load it 
    if file_path.endswith('.pdf'):
        loader = PyMuPDFLoader(file_path)
        documents = loader.load()
    elif file_path.endswith('.docx'):
        # Use python-docx directly to avoid docx2txt dependency
        doc = docx.Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        documents = [Document(page_content=full_text, metadata={"source": file_path, "page": 0})]
    else:
        raise ValueError("Unsupported file format! Please upload a PDF or DOCX file.")

    # 2. Split text into Chunks 
    # We chose RecursiveCharacterTextSplitter as it is best for preserving the context of legal paragraphs
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, 
        chunk_overlap=100
    )
    chunks = text_splitter.split_documents(documents)
    
    return chunks
